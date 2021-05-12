## This File auto matically play games and compute win rate
import subprocess
import os
import sys
from threading import Thread
from docx import Document
import time

# config --------------------------------------------------------------------------------------
PLAY_TIMES = 10   # Match Times
UPPER_LIST = ['RL_base'] 
LOWER_LIST = ['RL_base', 'IG', 'GreedyEnemy', 'IG2', 'RandomEnemy'] 
SHOW_BUG_MATCH = True  # show match detail of a bugged match
SHOW_TIMEOUT_MATCH = False  # show match detail of a timed out match
# config --------------------------------------------------------------------------------------





# A testing program
def test_program(PLAY_TIMES, UPPER, LOWER, doc_out):
    upper_wins = 0
    lower_wins = 0
    invincible_draws = 0
    same_state_draws = 0
    timeout_draws = 0
    i = PLAY_TIMES
    while (i):
        #Call commandline argument without getting its output
        #with open(os.devnull, "w") as f:
            #convert output to string
        output = ""
        try:
            output = subprocess.check_output(['python', '-m', 'referee', UPPER, LOWER]).decode(sys.stdout.encoding)
        except:
            output = subprocess.check_output(['python3', '-m', 'referee', UPPER, LOWER]).decode(sys.stdout.encoding)

        # get only last line of output
        result = output.splitlines()[-1]

        # count wining times
        if "upper" in result:
            upper_wins += 1
        elif "lower" in result:
            lower_wins += 1
        elif "draw" in result:
            # if draw because of both side has invincible symbol
            if "both" in result:
                invincible_draws += 1
            elif "same" in result:
                same_state_draws += 1
            else:
                if SHOW_TIMEOUT_MATCH:
                    print(output + " \n\n^^^^^^^^^^^^\nTIMEOUT Match Occurs!, \nMatch Detail above\n^^^^^^^^^^^^")
                    return
                timeout_draws += 1
        else:
            if SHOW_BUG_MATCH:
                print(output + " \n\n^^^^^^^^^^^^\nBUG Occurs!, \nMatch Detail above\n^^^^^^^^^^^^")
            return
        i -= 1

    # print("\n=======================================\n")
    # print("Total Plays: ", PLAY_TIMES)
    # print("Upper:", UPPER)
    # print("Lower:", LOWER)
    # print("---------------------------------------------")
    # print("\Win' RATE: \n")
    # print("  "+UPPER + " WIN RATE: ", round(upper_wins/PLAY_TIMES,2))
    # print("  "+LOWER + " WIN RATE: ", round(lower_wins/PLAY_TIMES,2))
    # print("\nDRAW RATE: \n")
    # print("  Both Invincible Rate:", round(invincible_draws/PLAY_TIMES,2))
    # print("  Same State Rate:", round(same_state_draws/PLAY_TIMES,2))
    # print("  Timeout Rate:", round(timeout_draws/PLAY_TIMES,2))
    # print("\n=======================================\n")
    format_list = [
        "\n=======================================\n",
        "\n\n   \t  "+UPPER+" \t VS \t "+ " \t "+LOWER,
        "-------------",
        "\nWin' RATE: \n",
        "  "+UPPER + " : " + str(round(upper_wins/PLAY_TIMES,3)),
        "  "+LOWER + " : " + str(round(lower_wins/PLAY_TIMES,3)),
        "-------------",
        "\nDRAW RATE: \n",
        "  Both Invincible Rate: " + str(round(invincible_draws/PLAY_TIMES,3)),
        "  Same State Rate: "+ str(round(same_state_draws/PLAY_TIMES,3)),
        "  Timeout Rate: "+ str(round(timeout_draws/PLAY_TIMES,3)),(
        "\n=======================================\n\n")
    ]
    doc_out.append(format_list)

def get_current_time():
    current_time = time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time()))
    return current_time

# Run the program with thread

if __name__ == '__main__':
    threads = []
    doc_out=[]
    for upr in UPPER_LIST:
        for lwr in LOWER_LIST:
            print("Testing  "+ upr +" Against " + lwr+ " .......")
            process = Thread(target=test_program(PLAY_TIMES,upr,lwr, doc_out))
            process.start()
            threads.append(process)
    for process in threads:
        process.join()

    print("Saving to Doc File....")
    # Save to Doc
    doc = Document()
    doc.add_paragraph(get_current_time())
    doc.add_paragraph("\n\nRounds For Each Pair: " + str(PLAY_TIMES))
    doc.add_paragraph("\n\nResult: " + '\n\n')
    for item in doc_out:
        for i in item:
            doc.add_paragraph(str(i))
    doc.save("TEST_RESULT.docx")
    print("Done!")
    
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[0])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[1])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[2])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[3])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[4])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[5])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[6])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[7])).start()


    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[0],LOWER_LIST[0])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[1])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[2])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[3])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[4])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[5])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[6])).start()
    # Thread(target = test_program(PLAY_TIMES,UPPER_LIST[1],LOWER_LIST[7])).start()
